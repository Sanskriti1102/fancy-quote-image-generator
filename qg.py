import random
import json
from pathlib import Path
import requests
from bs4 import BeautifulSoup
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from docx import Document

def scrape_quotes():
    print("\n🔍 Scraping quotes from https://quotes.toscrape.com ...")
    url = 'https://quotes.toscrape.com/page/{}/'
    all_quotes = []
    page = 1
    while True:
        res = requests.get(url.format(page))
        soup = BeautifulSoup(res.text, 'html.parser')
        quotes = soup.select('.quote')
        if not quotes:
            break
        for q in quotes:
            text = q.find('span', class_='text').get_text(strip=True)
            author = q.find('small', class_='author').text
            full_quote = f"{text} — {author}"
            all_quotes.append(full_quote)
        page += 1
    with open("scraped_quotes.json", "w") as f:
        json.dump(all_quotes, f, indent=2)
    print(f"✅ Scraped and stored {len(all_quotes)} quotes!\n")
    return all_quotes

def load_scraped_quotes():
    path = Path("scraped_quotes.json")
    if path.exists():
        with open(path, 'r') as f:
            return json.load(f)
    return []

def generate_wordcloud(quotes):
    text = " ".join(quotes)
    wc = WordCloud(width=800, height=400, background_color='white').generate(text)
    wc.to_file("wordcloud.png")
    plt.imshow(wc, interpolation='bilinear')
    plt.axis('off')
    plt.show()
    keywords = list(set(text.lower().replace('—', '').replace('“', '').replace('”', '').split()))
    return random.sample(keywords, min(20, len(keywords)))  # Get top 20 keywords (or all available)

def style_quote(quote):
    return (
        "🔹 **Thought for Today**\n\n"
        "Every day, I come across words that don’t just fill space — they fill *perspective*.\n"
        "Here's one that made me pause today:\n\n"
        f"💬 \"{quote}\"\n\n"
        "Take a moment. Let it sink in. Then move forward with purpose. 🚀\n\n"
        "#DailyInspiration #MindsetMatters #Leadership #SelfGrowth #LinkedInPosts"
    )

def export_to_word(name, quote):
    doc = Document()
    doc.add_heading(f"LinkedIn Style Quote for {name}", level=1)
    doc.add_paragraph(quote)
    file_path = f"{name.replace(' ', '_')}_linkedin_quote.docx"
    doc.save(file_path)
    print(f"\n📄 Quote saved to: {file_path}")

def main():
    print("✨ LinkedIn Quote Post Generator ✨\n")
    name = input("Enter your name: ").strip()

    # Step 1: Scrape quotes from the website
    quotes = scrape_quotes()

    # Step 2: Generate word cloud from the scraped quotes
    keywords = generate_wordcloud(quotes)

    print("\n💡 Keywords from the word cloud:")
    print(", ".join(keywords))

    # Step 3: Ask user to pick keywords based on the word cloud
    user_keywords = input("\nType keywords you want to search for (comma-separated): ").strip().lower()
    keyword_list = [k.strip() for k in user_keywords.split(',') if k.strip()]

    # Step 4: Re-scrape quotes based on user-selected keywords
    matched_quotes = [q for q in quotes if any(k in q.lower() for k in keyword_list)]

    if not matched_quotes:
        print("❗ No matching quotes found for your keywords. Showing a random quote instead.")
        matched_quotes = [random.choice(quotes)]

    print(f"\n🔹 Found {len(matched_quotes)} matching quote(s):\n")
    for idx, quote in enumerate(matched_quotes, 1):
        print(f"{idx}. {quote}\n")

    # Step 5: Allow user to select a quote to style and export
    choice = input("Pick a quote number to style and export (default 1): ").strip()
    try:
        selected = matched_quotes[int(choice) - 1]
    except:
        selected = matched_quotes[0]

    styled = style_quote(selected)

    print("\n✨ Your LinkedIn-style post:\n")
    print(styled)

    export_to_word(name, styled)

if __name__ == "__main__":
    main()
